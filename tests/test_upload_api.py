import io
import json
from pathlib import Path

from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pytest import MonkeyPatch

from backend.api.routes import documents as document_routes
from backend.main import create_app
from backend.services.document_service import DocumentService


def _build_client(database_path: Path, monkeypatch: MonkeyPatch) -> TestClient:
    service = DocumentService(database_path=str(database_path))
    monkeypatch.setattr(document_routes, "document_service", service)
    return TestClient(create_app())


def _pdf_bytes(text: str) -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=400, height=400)
    buffer = io.BytesIO()
    writer.write(buffer)
    # Real text extraction needs content; use a minimal valid PDF and accept
    # that extract_text may return "" for blank pages — we assert on metadata.
    return buffer.getvalue()


def test_upload_text_file_extracts_and_persists(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    client = _build_client(tmp_path / "uploads.db", monkeypatch)

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("background.txt", b"We are an SME specialising in robotics inspection since 2015.", "text/plain")},
        data={"application_id": "doc-abc-001"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "background.txt"
    assert payload["contentType"] == "text/plain"
    assert payload["characterCount"] == len("We are an SME specialising in robotics inspection since 2015.")
    assert "robotics inspection" in payload["textSnippet"]
    assert payload["applicationId"] == "doc-abc-001"
    assert payload["id"].startswith("upload-")
    assert payload["uploadedAt"]

    stored = document_routes.document_service.list_uploads("doc-abc-001")
    assert len(stored) == 1
    uploads = document_routes.document_service.application_store.list_uploads_for_application("doc-abc-001")
    assert uploads[0]["textSnippet"] == payload["textSnippet"]


def test_upload_docx_file_extracts_paragraphs(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    client = _build_client(tmp_path / "uploads_docx.db", monkeypatch)

    import docx

    document = docx.Document()
    document.add_paragraph("Northlight Robotics builds autonomous inspection drones.")
    document.add_paragraph("Founded in Munich with 40 employees.")
    buffer = io.BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("company.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert "inspection drones" in payload["textSnippet"]
    assert "Munich" in payload["textSnippet"]


def test_upload_pdf_file_returns_metadata(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    client = _build_client(tmp_path / "uploads_pdf.db", monkeypatch)

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("report.pdf", _pdf_bytes("ignored"), "application/pdf")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "report.pdf"
    assert payload["contentType"] == "application/pdf"
    assert isinstance(payload["characterCount"], int)


def test_upload_rejects_unsupported_type(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    client = _build_client(tmp_path / "uploads_bad.db", monkeypatch)

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("photo.png", b"\x89PNG-not-really", "image/png")},
    )

    assert response.status_code == 422
    assert "Unsupported file type" in response.json()["detail"]


def test_upload_rejects_oversized_file(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    from backend.services import upload_service

    client = _build_client(tmp_path / "uploads_big.db", monkeypatch)
    original_limit = upload_service.MAX_UPLOAD_BYTES
    upload_service.MAX_UPLOAD_BYTES = 10

    try:
        response = client.post(
            "/api/v1/documents/upload",
            files={"file": ("big.txt", b"x" * 100, "text/plain")},
        )
    finally:
        upload_service.MAX_UPLOAD_BYTES = original_limit

    assert response.status_code == 422
    assert "upload limit" in response.json()["detail"]


def test_upload_sanitizes_control_characters(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    client = _build_client(tmp_path / "uploads_sanitize.db", monkeypatch)

    raw = json.dumps({"about": "line1\x00\x01line2   spaced\n\n\n\n\nend"}).encode("utf-8")
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("data.json", raw, "application/json")},
    )

    assert response.status_code == 200
    snippet = response.json()["textSnippet"]
    assert "\x00" not in snippet
    assert "\x01" not in snippet


def test_upload_without_application_id(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    client = _build_client(tmp_path / "uploads_free.db", monkeypatch)

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("notes.md", b"# Heading\nSome notes.", "text/markdown")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["applicationId"] is None
