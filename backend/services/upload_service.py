"""In-memory text extraction for uploaded grant documents.

Supported formats: PDF (.pdf), Word (.docx), and plain text formats
(.txt, .md, .csv, .json). No files are written to disk — everything is
extracted from bytes in memory.
"""

from __future__ import annotations

import io
import re

from backend.core.logging import get_logger

logger = get_logger("services.upload")

MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_STORED_CHARACTERS = 200_000
SNIPPET_PREVIEW_LENGTH = 500

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json"}

SUPPORTED_SUFFIXES = TEXT_SUFFIXES | {".pdf", ".docx"}


class UnsupportedUploadError(ValueError):
    pass


class UploadTooLargeError(ValueError):
    pass


def _sanitize(text: str) -> str:
    """Strip control characters (except newlines/tabs) and collapse excess whitespace."""
    text = re.sub(r"[^\S\n]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(filename: str, content: bytes) -> str:
    """Extract plain text from raw upload bytes based on the file extension."""
    if len(content) > MAX_UPLOAD_BYTES:
        raise UploadTooLargeError(f"File exceeds the {MAX_UPLOAD_BYTES // (1024 * 1024)} MB upload limit.")

    suffix = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
    if suffix not in SUPPORTED_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise UnsupportedUploadError(f"Unsupported file type '{suffix or filename}'. Supported: {supported}.")

    try:
        if suffix == ".pdf":
            text = _extract_pdf(content)
        elif suffix == ".docx":
            text = _extract_docx(content)
        else:
            text = content.decode("utf-8", errors="replace")
    except (UnsupportedUploadError, UploadTooLargeError):
        raise
    except Exception as exc:
        logger.warning("Text extraction failed for '%s': %s", filename, exc)
        raise ValueError(f"Could not read '{filename}'. The file may be corrupted or password-protected.") from exc

    sanitized = _sanitize(text)
    return sanitized[:MAX_STORED_CHARACTERS]


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def detect_content_type(filename: str) -> str:
    suffix = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".csv": "text/csv",
        ".json": "application/json",
    }.get(suffix, "application/octet-stream")
