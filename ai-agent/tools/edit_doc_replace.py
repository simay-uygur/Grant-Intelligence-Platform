# tools/edit_doc_replace.py
# Stage 3 tool: find and replace text/sections in a Word (.docx) grant application.
# Used when the user asks to change part of a drafted document.

from docx import Document


def edit_doc_replace(doc_path, old_text, new_text, output_path=None):
    """
    Replace occurrences of old_text with new_text in a .docx file.

    doc_path:     path to the Word document to edit
    old_text:     the text/section to find
    new_text:     what to replace it with
    output_path:  where to save (defaults to overwriting doc_path)

    Returns the number of paragraphs where a replacement happened.
    """
    document = Document(doc_path)
    replacements = 0

    # Go through every paragraph and replace the text if found.
    for paragraph in document.paragraphs:
        if old_text in paragraph.text:
            # Rebuild the paragraph text with the replacement.
            new_paragraph_text = paragraph.text.replace(old_text, new_text)

            # Clear the paragraph's existing runs, then write the new text.
            # (A "run" is a piece of text with its own formatting.)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_paragraph_text
            else:
                paragraph.add_run(new_paragraph_text)

            replacements += 1

    # Also check inside tables (grant forms often use tables).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            new_p = paragraph.text.replace(old_text, new_text)
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = new_p
                            else:
                                paragraph.add_run(new_p)
                            replacements += 1

    # Save the edited document.
    save_path = output_path or doc_path
    document.save(save_path)

    print(f"[edit_doc_replace] Replaced text in {replacements} place(s). Saved to {save_path}")
    return replacements
