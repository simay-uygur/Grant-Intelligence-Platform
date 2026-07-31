# tests/make_test_doc.py
# Creates a simple .docx grant application draft to test edit_doc_replace against.

from docx import Document

doc = Document()
doc.add_heading("Grant Application Draft", level=1)
doc.add_paragraph("Organization: Green Energy Solutions Ltd")
doc.add_paragraph("Project goal: Develop solar panel recycling technology")
doc.add_paragraph("Requested budget: 250000 EUR")
doc.add_paragraph(
    "Summary: This project aims to build a scalable process for recycling "
    "end-of-life solar panels across the EU."
)

doc.save("tests/sample_application.docx")
print("Test document created: tests/sample_application.docx")